"""
Contract AI Service - AI-Powered Contract Intelligence
Uses Emergent LLM Integration for contract analysis, extraction, and advisory
"""
import os
import json
import re
import logging
from typing import Optional, Dict, Any, List
from datetime import datetime, timezone
from dotenv import load_dotenv

load_dotenv()

logger = logging.getLogger(__name__)

# Import models
from models.contract_governance import (
    ContractClassification,
    ContractRiskLevel,
    ContractAIExtraction,
    ContractAIAdvisory,
    DraftingHint,
    ClauseSuggestion,
    ConsistencyWarning,
    ContractRiskAssessment,
    ContractDDAnalysis,
    SERVICE_AGREEMENT_EXHIBITS,
)


# ==================== SYSTEM PROMPTS ====================

CONTRACT_EXTRACTION_PROMPT = """You are a Contract Document Analyst specializing in service agreements.
Extract key information from the contract document provided.

Extract the following:
1. Statement of Work (SOW) - Summary and key deliverables
2. Service Level Agreement (SLA) - Response times, priority levels, metrics
3. Contract Duration - Start date, end date, duration in months
4. Contract Value - Total value, currency, payment milestones
5. Parties - Supplier name, supplier country
6. Key Exhibits identified in the document

Return as JSON:
{
    "sow_summary": "Brief summary of scope...",
    "sow_details": "Detailed scope description...",
    "sla_summary": "SLA summary...",
    "sla_details": [{"priority": "Critical", "response_time": "...", "resolution_time": "..."}],
    "extracted_start_date": "YYYY-MM-DD or null",
    "extracted_end_date": "YYYY-MM-DD or null",
    "extracted_duration_months": number or null,
    "extracted_value": number or null,
    "extracted_currency": "SAR/USD/EUR",
    "extracted_milestones": [{"name": "...", "percentage": "...", "amount": "..."}],
    "supplier_name": "...",
    "supplier_country": "...",
    "exhibits_identified": ["Exhibit 1 - Definitions", "Exhibit 2 - SOW", ...],
    "extraction_confidence": 0.0-1.0,
    "extraction_notes": "Any issues or uncertainties..."
}"""


CONTRACT_CLASSIFICATION_PROMPT = """You are a Contract Classification Specialist for a Saudi Arabian financial institution.
Classify the contract based on SAMA (Saudi Central Bank) outsourcing regulations.

Classification Categories:
1. NOT_OUTSOURCING - Standard procurement, no outsourcing characteristics
2. OUTSOURCING - Service performed on continuing basis that could be undertaken by the bank
3. MATERIAL_OUTSOURCING - Outsourcing with material impact if disrupted (requires SAMA NOC)
4. CLOUD_COMPUTING - Cloud-hosted services (requires SAMA notification)
5. INSOURCING - Work performed on bank premises by external staff
6. EXEMPTED - Market data providers, clearing/settlement, correspondent banking, utilities

Classification Indicators:
- Cloud services, data hosting outside KSA → CLOUD_COMPUTING
- Vendor operating services on behalf of company → OUTSOURCING
- Material business impact if disrupted → MATERIAL_OUTSOURCING
- Simple goods/products purchase → NOT_OUTSOURCING
- Staff augmentation on-site → INSOURCING

Based on the contract context and details provided, return:
{
    "classification": "NOT_OUTSOURCING|OUTSOURCING|MATERIAL_OUTSOURCING|CLOUD_COMPUTING|INSOURCING|EXEMPTED",
    "classification_reason": "Detailed explanation...",
    "confidence": 0.0-1.0,
    "indicators_found": ["indicator1", "indicator2"],
    "requires_sama_noc": true/false,
    "requires_contract_dd": true/false
}"""


CONTRACT_ADVISORY_PROMPT = """You are a Contract Governance Advisor for a Saudi Arabian financial institution.
Provide drafting hints and clause suggestions based on the contract classification and context.

Service Agreement Exhibits:
1. Definitions
2. Statement of Work
3. Service Levels
4. Pricing and Financial Provisions
5. Human Resources
6. Sites
7. Technical Architecture
8. Reports
9. Business Continuity Plan and Disaster Recovery Plan
10. Duration
11. Supplier Insurance Requirements
12. Termination Assistance Compensation (Material Outsourcing only)
13. Termination Assistance (Material Outsourcing only)
14. Data Processing Agreement (Personal Data only)

Provide:
1. Drafting hints for relevant exhibits
2. Clause suggestions if outsourcing/cloud is detected
3. Consistency warnings comparing PR scope vs Contract scope

Return as JSON:
{
    "drafting_hints": [
        {"exhibit_number": 1, "exhibit_name": "Definitions", "hint_text": "...", "is_critical": false, "relevant_for": ["all"]}
    ],
    "clause_suggestions": [
        {"clause_type": "exit_plan", "clause_title": "Exit Strategy", "clause_text": "...", "reason": "...", "exhibit_reference": "Exhibit 13", "is_mandatory": true}
    ],
    "consistency_warnings": [
        {"warning_type": "scope_mismatch", "severity": "medium", "pr_value": "...", "contract_value": "...", "description": "..."}
    ],
    "ai_analysis_notes": "Overall observations..."
}"""


CONTRACT_DD_ANALYSIS_PROMPT = """You are a Contract Due Diligence Analyst for a Saudi Arabian financial institution.
Analyze the Due Diligence questionnaire responses and provide a risk assessment.

Sections to analyze:
1. Business Continuity - BCP, DRP, alternative sites, staff training
2. Anti-Fraud - Internal fraud history, theft incidents
3. Operational Risks - Licenses, outside KSA services, complaints
4. Cyber Security - Cloud, data outside KSA, digital channels, third-party access
5. Safety & Security - Physical security, 24/7 coverage, equipment
6. Judicial/Legal - Legal representation requirements
7. Regulatory - Government regulation, audited financials
8. Data Management - Data protection policies
9. SAMA Compliance - Consumer protection understanding

Risk Scoring:
- Positive answers to protective measures = lower risk
- Negative answers to red flag questions = higher risk
- Missing/unclear information = moderate risk increase

Return as JSON:
{
    "dd_risk_level": "Low|Medium|High",
    "dd_risk_score": 0-100,
    "key_findings": ["finding1", "finding2"],
    "missing_items": ["item1", "item2"],
    "required_followups": ["followup1", "followup2"],
    "business_continuity_summary": "...",
    "anti_fraud_summary": "...",
    "operational_risks_summary": "...",
    "cyber_security_summary": "...",
    "safety_security_summary": "...",
    "regulatory_summary": "...",
    "data_management_summary": "...",
    "sama_compliance_summary": "...",
    "analysis_confidence": 0.0-1.0
}"""


class ContractAIService:
    """AI Service for Contract Intelligence using OpenAI"""
    
    def __init__(self):
        """Initialize with OpenAI API key"""
        self.openai_key = os.environ.get("OPENAI_API_KEY")
        if not self.openai_key:
            logger.warning("No OPENAI_API_KEY provided. AI features will be disabled.")
        self.client = None
        if self.openai_key:
            from openai import OpenAI
            self.client = OpenAI(api_key=self.openai_key)
    
    async def extract_contract_document(self, file_path: str, file_type: str) -> str:
        """Extract text from uploaded contract document"""
        try:
            if file_type.lower() == "pdf":
                return await self._extract_pdf_text(file_path)
            elif file_type.lower() in ["docx", "doc"]:
                return await self._extract_docx_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error extracting contract document: {e}")
            raise
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF"""
        import subprocess
        try:
            result = subprocess.run(
                ["pdftotext", "-layout", file_path, "-"],
                capture_output=True,
                text=True,
                timeout=60
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout
        except Exception as e:
            logger.warning(f"pdftotext failed: {e}")
        return "[PDF requires OCR - text extraction not available]"
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from Word document"""
        try:
            from docx import Document
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    full_text.append(" | ".join(row_text))
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            return f"[Error extracting DOCX: {str(e)}]"
    
    async def extract_contract_fields(self, document_text: str) -> ContractAIExtraction:
        """Extract structured fields from contract document using AI"""
        
        if not self.client:
            raise ValueError("OPENAI_API_KEY required for contract extraction")
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": CONTRACT_EXTRACTION_PROMPT},
                    {"role": "user", "content": f"Extract information from this contract document:\n\n{document_text[:20000]}"}
                ],
                temperature=0.1
            )
            
            result_text = response.choices[0].message.content
            
            try:
                json_match = re.search(r'\{[\s\S]*\}', result_text)
                if json_match:
                    data = json.loads(json_match.group())
                    return ContractAIExtraction(
                        sow_summary=data.get("sow_summary"),
                        sow_details=data.get("sow_details"),
                        sla_summary=data.get("sla_summary"),
                        sla_details=data.get("sla_details", []),
                        extracted_start_date=data.get("extracted_start_date"),
                        extracted_end_date=data.get("extracted_end_date"),
                        extracted_duration_months=data.get("extracted_duration_months"),
                        extracted_value=data.get("extracted_value"),
                        extracted_currency=data.get("extracted_currency", "SAR"),
                        extracted_milestones=data.get("extracted_milestones", []),
                        supplier_name=data.get("supplier_name"),
                        supplier_country=data.get("supplier_country"),
                        exhibits_identified=data.get("exhibits_identified", []),
                        extraction_confidence=data.get("extraction_confidence"),
                        extraction_notes=data.get("extraction_notes"),
                        extracted_at=datetime.now(timezone.utc)
                    )
            except json.JSONDecodeError as e:
                logger.error(f"JSON decode error: {e}")
            
            return ContractAIExtraction(extraction_notes="AI extraction failed")
            
        except Exception as e:
            logger.error(f"Contract extraction failed: {str(e)}")
            raise
    
    async def classify_contract(
        self, 
        context_questionnaire: Dict[str, Any],
        contract_details: Dict[str, Any],
        vendor_info: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Classify contract based on SAMA outsourcing regulations"""
        
        if not self.client:
            # Fallback to rule-based classification
            return self._rule_based_classification(context_questionnaire, contract_details)
        
        try:
            context = f"""
CONTRACT CONTEXT QUESTIONNAIRE:
{json.dumps(context_questionnaire, indent=2, default=str)}

CONTRACT DETAILS:
{json.dumps(contract_details, indent=2, default=str)}

VENDOR INFORMATION:
{json.dumps(vendor_info, indent=2, default=str) if vendor_info else "Not provided"}
"""
            
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": CONTRACT_CLASSIFICATION_PROMPT},
                    {"role": "user", "content": f"Classify this contract:\n{context}"}
                ],
                temperature=0.1
            )
            
            result_text = response.choices[0].message.content
            
            try:
                json_match = re.search(r'\{[\s\S]*\}', result_text)
                if json_match:
                    return json.loads(json_match.group())
            except json.JSONDecodeError:
                pass
            
            return self._rule_based_classification(context_questionnaire, contract_details)
            
        except Exception as e:
            logger.error(f"Contract classification failed: {str(e)}")
            return self._rule_based_classification(context_questionnaire, contract_details)
    
    def _rule_based_classification(
        self,
        context: Dict[str, Any],
        contract: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Fallback rule-based classification"""
        classification = ContractClassification.NOT_OUTSOURCING
        reasons = []
        requires_sama = False
        requires_dd = False
        
        # Check cloud indicators
        if context.get("is_cloud_based") == "yes":
            classification = ContractClassification.CLOUD_COMPUTING
            reasons.append("Cloud-based service detected")
            requires_sama = True
            requires_dd = True
        
        # Check outsourcing indicators
        elif context.get("is_outsourcing_service") == "yes":
            classification = ContractClassification.OUTSOURCING
            reasons.append("Vendor operates service on behalf of company")
            requires_dd = True
            
            # Check materiality
            if context.get("expected_data_location") == "outside_ksa":
                classification = ContractClassification.MATERIAL_OUTSOURCING
                reasons.append("Data processed outside KSA")
                requires_sama = True
        
        # Check data access
        elif context.get("requires_system_data_access") == "yes":
            classification = ContractClassification.OUTSOURCING
            reasons.append("Service requires access to company systems/data")
            requires_dd = True
        
        return {
            "classification": classification.value,
            "classification_reason": "; ".join(reasons) if reasons else "Standard procurement - no outsourcing indicators",
            "confidence": 0.8,
            "indicators_found": reasons,
            "requires_sama_noc": requires_sama,
            "requires_contract_dd": requires_dd
        }
    
    async def generate_advisory(
        self,
        classification: str,
        context_questionnaire: Dict[str, Any],
        contract_details: Dict[str, Any],
        pr_details: Optional[Dict[str, Any]] = None
    ) -> ContractAIAdvisory:
        """Generate AI advisory including drafting hints and clause suggestions"""
        
        # Generate base drafting hints
        drafting_hints = self._generate_base_drafting_hints(classification)
        clause_suggestions = []
        consistency_warnings = []
        
        # Generate clause suggestions based on classification
        if classification in ["outsourcing", "material_outsourcing", "cloud_computing"]:
            clause_suggestions = self._generate_clause_suggestions(classification)
        
        # Check consistency with PR if provided
        if pr_details:
            consistency_warnings = self._check_pr_consistency(pr_details, contract_details)
        
        advisory = ContractAIAdvisory(
            drafting_hints=drafting_hints,
            clause_suggestions=clause_suggestions,
            consistency_warnings=consistency_warnings,
            generated_at=datetime.now(timezone.utc),
            last_updated=datetime.now(timezone.utc)
        )
        
        # Try to enhance with AI
        if self.client:
            try:
                context = f"""
CLASSIFICATION: {classification}

CONTEXT QUESTIONNAIRE:
{json.dumps(context_questionnaire, indent=2, default=str)}

CONTRACT DETAILS:
{json.dumps(contract_details, indent=2, default=str)}

PR DETAILS:
{json.dumps(pr_details, indent=2, default=str) if pr_details else "Not provided"}
"""
                
                response = self.client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": CONTRACT_ADVISORY_PROMPT},
                        {"role": "user", "content": f"Generate advisory for this contract:\n{context}"}
                    ],
                    temperature=0.1
                )
                
                result_text = response.choices[0].message.content
                
                try:
                    json_match = re.search(r'\{[\s\S]*\}', result_text)
                    if json_match:
                        data = json.loads(json_match.group())
                        advisory.ai_analysis_notes = data.get("ai_analysis_notes")
                except:
                    pass
                    
            except Exception as e:
                logger.warning(f"AI advisory enhancement failed: {e}")
        
        return advisory
    
    def _generate_base_drafting_hints(self, classification: str) -> List[DraftingHint]:
        """Generate base drafting hints for exhibits"""
        hints = []
        
        for exhibit in SERVICE_AGREEMENT_EXHIBITS:
            hint_text = f"Review {exhibit['name']} carefully."
            is_critical = exhibit.get("always_required", False)
            relevant_for = ["all"]
            
            # Add classification-specific hints
            if exhibit["number"] == 2:  # SOW
                hint_text = "Ensure scope is clearly defined with measurable deliverables. Include acceptance criteria."
            elif exhibit["number"] == 3:  # SLA
                hint_text = "Define priority levels, response times, and resolution times. Include penalties for SLA breaches."
            elif exhibit["number"] == 9:  # BCP/DRP
                if classification in ["outsourcing", "material_outsourcing", "cloud_computing"]:
                    hint_text = "CRITICAL: For outsourcing/cloud services, ensure BCP and DRP are comprehensive. Include disaster recovery site, rehearsal schedules, and recovery time objectives."
                    is_critical = True
                    relevant_for = ["outsourcing", "cloud"]
            elif exhibit["number"] == 12:  # Termination Assistance Compensation
                if classification == "material_outsourcing":
                    hint_text = "MANDATORY for Material Outsourcing: Define termination assistance compensation structure."
                    is_critical = True
                    relevant_for = ["material_outsourcing"]
            elif exhibit["number"] == 13:  # Termination Assistance
                if classification == "material_outsourcing":
                    hint_text = "MANDATORY for Material Outsourcing: Define exit plan, knowledge transfer, and transition support."
                    is_critical = True
                    relevant_for = ["material_outsourcing"]
            elif exhibit["number"] == 14:  # DPA
                if classification in ["outsourcing", "cloud_computing"]:
                    hint_text = "CRITICAL: Include Data Processing Agreement if personal data is involved. Comply with KSA PDPL."
                    is_critical = True
                    relevant_for = ["data_processing", "cloud"]
            
            hints.append(DraftingHint(
                exhibit_number=exhibit["number"],
                exhibit_name=exhibit["name"],
                hint_text=hint_text,
                is_critical=is_critical,
                relevant_for=relevant_for
            ))
        
        return hints
    
    def _generate_clause_suggestions(self, classification: str) -> List[ClauseSuggestion]:
        """Generate clause suggestions based on classification"""
        suggestions = []
        
        if classification in ["outsourcing", "material_outsourcing"]:
            suggestions.append(ClauseSuggestion(
                clause_type="audit_rights",
                clause_title="Audit Rights Clause",
                clause_text="The Bank reserves the right to audit the Supplier's facilities, systems, and processes related to the Services at any time with reasonable notice.",
                reason="Required for outsourcing contracts to ensure compliance and oversight",
                exhibit_reference="Master Services Agreement",
                is_mandatory=True
            ))
            suggestions.append(ClauseSuggestion(
                clause_type="subcontracting",
                clause_title="Subcontracting Restrictions",
                clause_text="The Supplier shall not subcontract any part of the Services without prior written approval from the Bank.",
                reason="Control over service delivery chain",
                exhibit_reference="Master Services Agreement",
                is_mandatory=True
            ))
        
        if classification == "material_outsourcing":
            suggestions.append(ClauseSuggestion(
                clause_type="exit_plan",
                clause_title="Exit Strategy and Termination Assistance",
                clause_text="Upon termination, the Supplier shall provide termination assistance for a period of [X] months to ensure smooth transition of services.",
                reason="SAMA requirement for material outsourcing",
                exhibit_reference="Exhibit 13",
                is_mandatory=True
            ))
            suggestions.append(ClauseSuggestion(
                clause_type="sama_notification",
                clause_title="SAMA Notification Clause",
                clause_text="This contract is subject to SAMA notification requirements. The Bank shall notify SAMA of this material outsourcing arrangement.",
                reason="Regulatory compliance for material outsourcing",
                is_mandatory=True
            ))
        
        if classification == "cloud_computing":
            suggestions.append(ClauseSuggestion(
                clause_type="data_location",
                clause_title="Data Location and Sovereignty",
                clause_text="All data processing and storage shall occur within the Kingdom of Saudi Arabia unless explicitly approved by the Bank.",
                reason="Data sovereignty requirement for cloud services",
                exhibit_reference="Exhibit 14",
                is_mandatory=True
            ))
            suggestions.append(ClauseSuggestion(
                clause_type="security_standards",
                clause_title="Cloud Security Standards",
                clause_text="The Supplier shall maintain ISO 27001 certification and comply with SAMA Cyber Security Framework.",
                reason="Security requirements for cloud services",
                is_mandatory=True
            ))
        
        return suggestions
    
    def _check_pr_consistency(self, pr_details: Dict[str, Any], contract_details: Dict[str, Any]) -> List[ConsistencyWarning]:
        """Check consistency between PR and Contract"""
        warnings = []
        
        # Check value mismatch - ensure numeric conversion
        try:
            pr_budget = float(pr_details.get("budget", 0) or 0)
            contract_value = float(contract_details.get("value", 0) or 0)
        except (TypeError, ValueError):
            pr_budget = 0
            contract_value = 0
            
        if pr_budget and contract_value:
            variance = abs(contract_value - pr_budget) / pr_budget * 100 if pr_budget > 0 else 0
            if variance > 20:
                warnings.append(ConsistencyWarning(
                    warning_type="value_mismatch",
                    severity="high" if variance > 50 else "medium",
                    pr_value=str(pr_budget),
                    contract_value=str(contract_value),
                    description=f"Contract value differs from PR budget by {variance:.1f}%"
                ))
        
        # Check scope mismatch
        pr_scope = pr_details.get("requirements", "")
        contract_scope = contract_details.get("sow", "")
        if pr_scope and contract_scope and len(pr_scope) > 50 and len(contract_scope) > 50:
            # Simple check - in production, use semantic similarity
            if len(contract_scope) < len(pr_scope) * 0.5:
                warnings.append(ConsistencyWarning(
                    warning_type="scope_mismatch",
                    severity="medium",
                    pr_value=pr_scope[:100] + "...",
                    contract_value=contract_scope[:100] + "...",
                    description="Contract scope appears significantly shorter than PR requirements"
                ))
        
        return warnings
    
    async def analyze_contract_dd(
        self,
        dd_responses: List[Dict[str, Any]],
        document_text: Optional[str] = None
    ) -> ContractDDAnalysis:
        """Analyze Contract Due Diligence questionnaire responses"""
        
        # Calculate rule-based scores first
        analysis = self._rule_based_dd_analysis(dd_responses)
        
        # Enhance with AI if available
        if self.client:
            try:
                context = f"""
DUE DILIGENCE RESPONSES:
{json.dumps(dd_responses, indent=2, default=str)}

DOCUMENT TEXT (if available):
{document_text[:10000] if document_text else "Not provided"}
"""
                
                response = self.client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {"role": "system", "content": CONTRACT_DD_ANALYSIS_PROMPT},
                        {"role": "user", "content": f"Analyze this Due Diligence:\n{context}"}
                    ],
                    temperature=0.1
                )
                
                result_text = response.choices[0].message.content
                
                try:
                    json_match = re.search(r'\{[\s\S]*\}', result_text)
                    if json_match:
                        data = json.loads(json_match.group())
                        # Update analysis with AI insights
                        analysis.key_findings = data.get("key_findings", analysis.key_findings)
                        analysis.missing_items = data.get("missing_items", analysis.missing_items)
                        analysis.required_followups = data.get("required_followups", analysis.required_followups)
                        analysis.business_continuity_summary = data.get("business_continuity_summary")
                        analysis.cyber_security_summary = data.get("cyber_security_summary")
                        analysis.analysis_confidence = data.get("analysis_confidence", 0.7)
                except:
                    pass
                    
            except Exception as e:
                logger.warning(f"AI DD analysis failed: {e}")
        
        analysis.analyzed_at = datetime.now(timezone.utc)
        return analysis
    
    def _rule_based_dd_analysis(self, dd_responses: List[Dict[str, Any]]) -> ContractDDAnalysis:
        """Rule-based DD analysis"""
        score = 50  # Start at medium
        findings = []
        missing = []
        followups = []
        
        # Group responses by section
        sections = {}
        for resp in dd_responses:
            section = resp.get("section", "Unknown")
            if section not in sections:
                sections[section] = []
            sections[section].append(resp)
        
        # Analyze Business Continuity
        bc_responses = sections.get("Business Continuity", [])
        bc_yes_count = sum(1 for r in bc_responses if r.get("answer") == "yes")
        bc_no_count = sum(1 for r in bc_responses if r.get("answer") == "no")
        
        if len(bc_responses) > 0:
            bc_score = (bc_yes_count / len(bc_responses)) * 100
            if bc_score < 50:
                score += 15
                findings.append("Business Continuity practices need improvement")
                followups.append("Request BCP/DRP documentation")
        
        # Analyze Cyber Security
        cs_responses = sections.get("Cyber Security", [])
        # Red flags
        for resp in cs_responses:
            if resp.get("question_id") == "cs_2" and resp.get("answer") == "yes":
                score += 10
                findings.append("Data processed outside KSA")
                followups.append("Verify data sovereignty compliance")
            if resp.get("question_id") == "cs_6" and resp.get("answer") == "yes":
                score += 5
                findings.append("Third-party system access required")
        
        # Analyze Anti-Fraud
        af_responses = sections.get("Anti-Fraud", [])
        for resp in af_responses:
            if resp.get("answer") == "yes":
                score += 15
                findings.append(f"Fraud/theft incident reported: {resp.get('question_text', '')}")
        
        # Determine risk level
        if score < 40:
            risk_level = ContractRiskLevel.LOW
        elif score < 70:
            risk_level = ContractRiskLevel.MEDIUM
        else:
            risk_level = ContractRiskLevel.HIGH
        
        return ContractDDAnalysis(
            dd_risk_level=risk_level,
            dd_risk_score=min(score, 100),
            key_findings=findings,
            missing_items=missing,
            required_followups=followups,
            analysis_confidence=0.6
        )
    
    def calculate_contract_risk(
        self,
        classification: str,
        vendor_risk_score: float,
        context_questionnaire: Dict[str, Any],
        contract_value: float,
        duration_months: int
    ) -> ContractRiskAssessment:
        """Calculate overall contract risk assessment"""
        
        # Base score from vendor risk (30% weight)
        vendor_contribution = vendor_risk_score * 0.3
        
        # Data exposure risk (25% weight)
        data_risk = 0
        if context_questionnaire.get("requires_system_data_access") == "yes":
            data_risk += 15
        if context_questionnaire.get("expected_data_location") == "outside_ksa":
            data_risk += 25
        data_contribution = data_risk * 0.25
        
        # Outsourcing/Cloud risk (25% weight)
        outsourcing_risk = 0
        if classification == "cloud_computing":
            outsourcing_risk = 60
        elif classification == "material_outsourcing":
            outsourcing_risk = 70
        elif classification == "outsourcing":
            outsourcing_risk = 50
        outsourcing_contribution = outsourcing_risk * 0.25
        
        # Duration/Dependency risk (10% weight)
        duration_risk = 0
        if duration_months > 36:
            duration_risk = 60
        elif duration_months > 24:
            duration_risk = 40
        elif duration_months > 12:
            duration_risk = 20
        duration_contribution = duration_risk * 0.1
        
        # Value risk (10% weight)
        value_risk = 0
        if contract_value > 10000000:  # > 10M
            value_risk = 60
        elif contract_value > 5000000:  # > 5M
            value_risk = 40
        elif contract_value > 1000000:  # > 1M
            value_risk = 20
        value_contribution = value_risk * 0.1
        
        # Total risk score
        total_score = (
            vendor_contribution +
            data_contribution +
            outsourcing_contribution +
            duration_contribution +
            value_contribution
        )
        
        # Determine risk level
        if total_score < 40:
            risk_level = ContractRiskLevel.LOW
        elif total_score < 70:
            risk_level = ContractRiskLevel.MEDIUM
        else:
            risk_level = ContractRiskLevel.HIGH
        
        # Determine required actions
        requires_dd = classification in ["outsourcing", "material_outsourcing", "cloud_computing"]
        requires_sama = classification in ["material_outsourcing", "cloud_computing"]
        requires_acceptance = risk_level == ContractRiskLevel.HIGH
        
        # Generate risk drivers
        risk_drivers = []
        if vendor_contribution > 20:
            risk_drivers.append(f"High vendor risk score ({vendor_risk_score:.0f})")
        if data_contribution > 10:
            risk_drivers.append("Data exposure concerns")
        if outsourcing_contribution > 15:
            risk_drivers.append(f"Contract classified as {classification}")
        if duration_contribution > 5:
            risk_drivers.append(f"Long contract duration ({duration_months} months)")
        if value_contribution > 5:
            risk_drivers.append(f"High contract value (${contract_value:,.0f})")
        
        return ContractRiskAssessment(
            risk_score=min(total_score, 100),
            risk_level=risk_level,
            top_risk_drivers=risk_drivers[:5],
            vendor_risk_contribution=vendor_contribution,
            data_exposure_risk=data_contribution,
            outsourcing_cloud_risk=outsourcing_contribution,
            duration_dependency_risk=duration_contribution,
            value_risk=value_contribution,
            requires_contract_dd=requires_dd,
            requires_sama_noc=requires_sama,
            requires_risk_acceptance=requires_acceptance,
            assessed_by="ai",
            assessed_at=datetime.now(timezone.utc)
        )


# Singleton instance
_ai_service: Optional[ContractAIService] = None


def get_contract_ai_service() -> ContractAIService:
    """Get or create the AI service singleton"""
    global _ai_service
    if _ai_service is None:
        _ai_service = ContractAIService()
    return _ai_service
