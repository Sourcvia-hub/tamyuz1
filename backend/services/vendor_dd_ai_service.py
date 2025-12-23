"""
Vendor Due Diligence AI Service - Document extraction and risk evaluation
Uses Emergent LLM Integration for AI analysis
"""
import os
import json
import re
import logging
import asyncio
from typing import Optional, Dict, Any, Tuple
from datetime import datetime, timezone
from dotenv import load_dotenv

load_dotenv()

logger = logging.getLogger(__name__)

# AI System Prompt for Vendor Risk Evaluation
VENDOR_DD_SYSTEM_PROMPT = """You are a Vendor Due Diligence Risk Analysis Assistant operating in a regulated procurement and governance environment.
Your role is to analyze Vendor Due Diligence documents (Word or PDF) and produce a balanced, commercially realistic vendor risk assessment.
You must:
• Be practical and non-punitive
• Avoid idealistic or overly strict judgments
• Distinguish clearly between structural risks and maturity gaps
• Provide transparent reasoning suitable for audit, compliance, and management review
You do NOT approve, reject, or blacklist vendors.
You ONLY analyze and explain inherent vendor risk.
Final decisions are always made by human officers.

📥 INPUT
You will receive:
• A Vendor Due Diligence document (Word or PDF)
• The document may be incomplete, partially filled, or written in general language
• Missing or unclear information is expected and must not be treated as failure

🧠 RISK EVALUATION PRINCIPLES
1️⃣ Risk Weighting Logic (MANDATORY)
You must classify all findings into the following categories:

🔴 A. Structural / High-Impact Risks (HIGH WEIGHT)
These materially increase inherent vendor risk:
• Vendor headquartered or operating in high-risk jurisdictions
(e.g. Russia, sanctioned countries, conflict-affected regions)
• Exposure to international or local sanctions
• Ongoing or recent legal or regulatory actions
• Opaque or unclear ownership / beneficial ownership
These risks may result in High Risk even if documentation is otherwise strong.

🟠 B. Maturity & Governance Gaps (MEDIUM WEIGHT)
These are common in real markets and must be treated realistically:
• Absence of formal written policies
• Early-stage governance frameworks
• Generic or high-level responses
• Limited financial history or scale
These increase risk moderately only
They must NOT alone result in High Risk.

🟢 C. Missing or Unclear Information (LOW WEIGHT)
• Blank answers
• "Will be provided later"
• Ambiguous or unclear wording
These:
• Reduce confidence
• Trigger follow-up
• Must NOT significantly increase risk score

⚠️ SANCTIONS HANDLING RULE (MANDATORY)
Sanctions status must NOT be conclusively determined by you unless it is explicitly and clearly confirmed in the provided document.
If sanctions information is missing, unclear, or not verifiable:
• Do NOT conclude that the vendor is sanctioned or clear
• Mark sanctions as "Pending Officer Review"
• Clearly state that sanctions screening requires officer confirmation
The final vendor risk assessment must NOT be considered concluded until sanctions status has been reviewed and updated by the responsible officer.
You may:
• Flag sanctions as a potential risk driver
• Reduce confidence level if sanctions information is incomplete
You must:
• Clearly document this dependency in Notes for Human Review

📊 SCORING & THRESHOLDS (STRICT)
You must calculate a Vendor Risk Score from 0 to 100.
Risk Thresholds:
• 0 – 39 → Low Risk
• 40 – 69 → Medium Risk
• 70 – 100 → High Risk
Threshold Interpretation Rules:
• High Risk does NOT mean rejection
→ It means explicit risk acceptance and senior approval are required
• Medium Risk is expected and acceptable for many vendors
• Low Risk does NOT mean no risk

⚠️ OVERRIDE RULES
Apply these regardless of numeric score:
• If vendor headquarters is in a High-Risk Jurisdiction, minimum risk level = High
• If sanctions exposure is mentioned or suspected, flag as a critical risk driver
• If ownership transparency is weak, risk level cannot be lower than Medium

🚫 RESTRICTIONS
• Do NOT approve or reject vendors
• Do NOT assume contract scope, outsourcing, or cloud usage
• Do NOT require perfect documentation
• Always apply commercial realism
• Always defer final judgment to humans

✅ FINAL PRINCIPLE
Your purpose is to support informed human decision-making."""

# Field extraction prompt
FIELD_EXTRACTION_PROMPT = """You are a document data extraction specialist. Extract the following fields from the vendor registration document.
For each field, provide:
1. The extracted value (if found)
2. The status: "Extracted" (directly found), "Inferred" (reasonably deduced), or "Not Provided" (not found)
3. Confidence level (0.0 to 1.0)

Return the data as a JSON object with the following structure:
{
    "vendor_name_arabic": {"value": "...", "status": "Extracted/Inferred/Not Provided", "confidence": 0.0-1.0},
    "vendor_name_english": {"value": "...", "status": "...", "confidence": ...},
    "commercial_name": {"value": "...", "status": "...", "confidence": ...},
    "entity_type": {"value": "...", "status": "...", "confidence": ...},
    "vat_number": {"value": "...", "status": "...", "confidence": ...},
    "unified_number": {"value": "...", "status": "...", "confidence": ...},
    "cr_number": {"value": "...", "status": "...", "confidence": ...},
    "cr_expiry_date": {"value": "...", "status": "...", "confidence": ...},
    "cr_country_city": {"value": "...", "status": "...", "confidence": ...},
    "license_number": {"value": "...", "status": "...", "confidence": ...},
    "license_expiry_date": {"value": "...", "status": "...", "confidence": ...},
    "activity_description": {"value": "...", "status": "...", "confidence": ...},
    "employees_total": {"value": "...", "status": "...", "confidence": ...},
    "employees_saudi": {"value": "...", "status": "...", "confidence": ...},
    "address_street": {"value": "...", "status": "...", "confidence": ...},
    "address_building": {"value": "...", "status": "...", "confidence": ...},
    "address_city": {"value": "...", "status": "...", "confidence": ...},
    "address_district": {"value": "...", "status": "...", "confidence": ...},
    "address_country": {"value": "...", "status": "...", "confidence": ...},
    "contact_mobile": {"value": "...", "status": "...", "confidence": ...},
    "contact_landline": {"value": "...", "status": "...", "confidence": ...},
    "contact_fax": {"value": "...", "status": "...", "confidence": ...},
    "contact_email": {"value": "...", "status": "...", "confidence": ...},
    "rep_name": {"value": "...", "status": "...", "confidence": ...},
    "rep_designation": {"value": "...", "status": "...", "confidence": ...},
    "rep_id_type": {"value": "...", "status": "...", "confidence": ...},
    "rep_id_number": {"value": "...", "status": "...", "confidence": ...},
    "rep_nationality": {"value": "...", "status": "...", "confidence": ...},
    "rep_mobile": {"value": "...", "status": "...", "confidence": ...},
    "rep_email": {"value": "...", "status": "...", "confidence": ...},
    "bank_account_name": {"value": "...", "status": "...", "confidence": ...},
    "bank_name": {"value": "...", "status": "...", "confidence": ...},
    "bank_branch": {"value": "...", "status": "...", "confidence": ...},
    "bank_country": {"value": "...", "status": "...", "confidence": ...},
    "iban": {"value": "...", "status": "...", "confidence": ...},
    "currency": {"value": "...", "status": "...", "confidence": ...},
    "swift_code": {"value": "...", "status": "...", "confidence": ...},
    "years_in_business": {"value": "...", "status": "...", "confidence": ...},
    "number_of_customers": {"value": "...", "status": "...", "confidence": ...},
    "number_of_branches": {"value": "...", "status": "...", "confidence": ...},
    "owners_managers": [{"name": "...", "nationality": "...", "id_number": "...", "ownership_percentage": "..."}]
}

Important:
- Extract all visible text from forms, tables, and handwritten sections if legible
- For Arabic text, provide both Arabic and English transliteration if possible
- For dates, convert to ISO format (YYYY-MM-DD)
- For Yes/No questions, convert to boolean (true/false) or null if unclear
- Mark confidence as low (< 0.5) for handwritten or unclear text"""


class VendorDDAIService:
    """AI Service for Vendor Due Diligence using Emergent LLM Integration"""
    
    def __init__(self):
        """Initialize with OpenAI API key"""
        self.openai_key = os.environ.get("OPENAI_API_KEY")
        if not self.openai_key:
            logger.warning("No OPENAI_API_KEY provided. AI features will be disabled.")
        self.client = None
        if self.openai_key:
            from openai import OpenAI
            self.client = OpenAI(api_key=self.openai_key)
        
        # High-risk countries list (configurable)
        self.high_risk_countries = self._load_high_risk_countries()
    
    def _load_high_risk_countries(self) -> list:
        """Load high-risk countries from database or use defaults"""
        return [
            "Russia", "North Korea", "Iran", "Syria", "Cuba", "Venezuela",
            "Myanmar", "Belarus", "Zimbabwe", "Sudan", "South Sudan",
            "Democratic Republic of Congo", "Central African Republic",
            "Libya", "Somalia", "Yemen", "Afghanistan", "Iraq"
        ]
    
    async def extract_document_text(self, file_path: str, file_type: str) -> str:
        """Extract text from uploaded document"""
        try:
            if file_type.lower() == "pdf":
                return await self._extract_pdf_text(file_path)
            elif file_type.lower() in ["docx", "doc"]:
                return await self._extract_docx_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error extracting document text: {e}")
            raise
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF using pdftotext"""
        import subprocess
        
        # Try pdftotext first (for text-based PDFs)
        try:
            result = subprocess.run(
                ["pdftotext", "-layout", file_path, "-"],
                capture_output=True,
                text=True,
                timeout=60
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout
        except Exception as e:
            logger.warning(f"pdftotext failed: {e}")
        
        # If pdftotext fails, return a message that we need OCR
        return "[PDF requires OCR - text extraction not available]"
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from Word document"""
        try:
            from docx import Document
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    full_text.append(" | ".join(row_text))
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            return f"[Error extracting DOCX: {str(e)}]"
    
    async def extract_fields(self, document_text: str) -> Dict[str, Any]:
        """Extract structured fields from document text using Emergent LLM"""
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        
        if not self.emergent_key:
            raise ValueError("EMERGENT_LLM_KEY required for field extraction")
        
        try:
            # Initialize chat with Emergent key
            chat = LlmChat(
                api_key=self.emergent_key,
                session_id=f"vendor-dd-extract-{datetime.now().timestamp()}",
                system_message=FIELD_EXTRACTION_PROMPT
            ).with_model("openai", "gpt-4o")
            
            # Create message
            user_message = UserMessage(
                text=f"Extract fields from this vendor registration document:\n\n{document_text[:15000]}"
            )
            
            # Send message and get response
            response = await chat.send_message(user_message)
            
            # Parse JSON response
            try:
                # Try to extract JSON from response
                json_match = re.search(r'\{[\s\S]*\}', response)
                if json_match:
                    return json.loads(json_match.group())
                else:
                    logger.error(f"No JSON found in response: {response[:500]}")
                    return {}
            except json.JSONDecodeError as e:
                logger.error(f"JSON decode error: {e}")
                return {}
                
        except Exception as e:
            logger.error(f"Field extraction failed: {str(e)}")
            raise Exception(f"Field extraction failed: {str(e)}")
    
    async def run_risk_assessment(self, document_text: str, extracted_fields: Dict[str, Any]) -> Dict[str, Any]:
        """Run AI risk assessment on vendor document using Emergent LLM"""
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        
        if not self.emergent_key:
            raise ValueError("EMERGENT_LLM_KEY required for risk assessment")
        
        # Prepare context for risk assessment
        context = f"""
EXTRACTED VENDOR INFORMATION:
{json.dumps(extracted_fields, indent=2, default=str)}

RAW DOCUMENT TEXT:
{document_text[:10000]}
"""
        
        try:
            # Initialize chat with Emergent key
            chat = LlmChat(
                api_key=self.emergent_key,
                session_id=f"vendor-dd-risk-{datetime.now().timestamp()}",
                system_message=VENDOR_DD_SYSTEM_PROMPT
            ).with_model("openai", "gpt-4o")
            
            # Create message
            user_message = UserMessage(
                text=f"""Analyze this vendor due diligence information and provide a risk assessment.

{context}

Respond in the following JSON format:
{{
    "vendor_name": "...",
    "country_jurisdiction": "...",
    "vendor_risk_score": 0-100,
    "vendor_risk_level": "Low/Medium/High",
    "top_risk_drivers": ["driver1", "driver2", "driver3"],
    "assessment_summary": "...",
    "ai_confidence_level": "High/Medium/Low",
    "ai_confidence_rationale": "...",
    "notes_for_human_review": "..."
}}"""
            )
            
            # Send message and get response
            response = await chat.send_message(user_message)
            
            # Parse JSON response
            try:
                json_match = re.search(r'\{[\s\S]*\}', response)
                if json_match:
                    assessment = json.loads(json_match.group())
                    # Apply override rules
                    assessment = self._apply_risk_overrides(assessment, extracted_fields)
                    return assessment
                else:
                    logger.error(f"No JSON found in risk response")
                    return self._default_assessment()
            except json.JSONDecodeError as e:
                logger.error(f"JSON decode error in risk assessment: {e}")
                return self._default_assessment()
                
        except Exception as e:
            logger.error(f"Risk assessment failed: {str(e)}")
            raise Exception(f"Risk assessment failed: {str(e)}")
    
    def _default_assessment(self) -> Dict[str, Any]:
        """Return a default assessment when AI fails"""
        return {
            "vendor_name": "Unknown",
            "country_jurisdiction": "Unknown",
            "vendor_risk_score": 50,
            "vendor_risk_level": "Medium",
            "top_risk_drivers": ["Unable to complete AI assessment"],
            "assessment_summary": "AI assessment could not be completed. Manual review required.",
            "ai_confidence_level": "Low",
            "ai_confidence_rationale": "AI processing error",
            "notes_for_human_review": "Please review all vendor information manually."
        }
    
    def _apply_risk_overrides(self, assessment: Dict[str, Any], extracted_fields: Dict[str, Any]) -> Dict[str, Any]:
        """Apply mandatory risk override rules"""
        
        # Get country from assessment or extracted fields
        country = assessment.get("country_jurisdiction") or ""
        if not country and extracted_fields.get("address_country"):
            country = extracted_fields["address_country"].get("value", "") if isinstance(extracted_fields["address_country"], dict) else str(extracted_fields["address_country"])
        
        # Override 1: High-risk country = minimum High risk
        if country and any(hrc.lower() in country.lower() for hrc in self.high_risk_countries):
            if assessment.get("vendor_risk_level") != "High":
                assessment["vendor_risk_level"] = "High"
                if assessment.get("vendor_risk_score", 0) < 70:
                    assessment["vendor_risk_score"] = 70
                assessment["notes_for_human_review"] = (
                    assessment.get("notes_for_human_review", "") + 
                    "\n\n⚠️ HIGH-RISK JURISDICTION OVERRIDE: Vendor is located in a high-risk country. Minimum risk level set to High."
                )
                risk_drivers = assessment.get("top_risk_drivers", [])
                if not any("jurisdiction" in str(d).lower() for d in risk_drivers):
                    risk_drivers.insert(0, f"High-risk jurisdiction: {country}")
                    assessment["top_risk_drivers"] = risk_drivers[:3]
        
        # Override 2: Check for sanctions mentions
        notes = str(assessment.get("notes_for_human_review", "")).lower()
        summary = str(assessment.get("assessment_summary", "")).lower()
        if "sanction" in notes or "sanction" in summary:
            risk_drivers = assessment.get("top_risk_drivers", [])
            if not any("sanction" in str(d).lower() for d in risk_drivers):
                risk_drivers.insert(0, "Potential sanctions exposure - requires officer verification")
                assessment["top_risk_drivers"] = risk_drivers[:3]
        
        # Override 3: Weak ownership transparency = minimum Medium
        owners = extracted_fields.get("owners_managers", [])
        if not owners or len(owners) == 0:
            if assessment.get("vendor_risk_level") == "Low":
                assessment["vendor_risk_level"] = "Medium"
                if assessment.get("vendor_risk_score", 0) < 40:
                    assessment["vendor_risk_score"] = 40
                assessment["notes_for_human_review"] = (
                    assessment.get("notes_for_human_review", "") + 
                    "\n\n⚠️ OWNERSHIP OVERRIDE: Limited ownership transparency. Risk level cannot be lower than Medium."
                )
        
        return assessment
    
    def calculate_risk_level(self, score: float) -> str:
        """Calculate risk level based on score thresholds"""
        if score < 40:
            return "Low"
        elif score < 70:
            return "Medium"
        else:
            return "High"


# Singleton instance
_ai_service: Optional[VendorDDAIService] = None

def get_vendor_dd_ai_service() -> VendorDDAIService:
    """Get or create the AI service singleton"""
    global _ai_service
    if _ai_service is None:
        _ai_service = VendorDDAIService()
    return _ai_service
